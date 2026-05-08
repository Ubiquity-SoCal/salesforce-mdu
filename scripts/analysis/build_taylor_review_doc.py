from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

# Title
title = doc.add_heading("Taylor's Salesforce Feedback — Review & Implementation Plan", level=1)
title.runs[0].font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

doc.add_paragraph("Source: Two emails from Taylor Mauney (2026-03-30)")
doc.add_paragraph("Status: Pending review and implementation")
doc.add_paragraph("")

# ============================================================
# SECTION: Assessment Summary
# ============================================================
doc.add_heading("Assessment Summary", level=2)

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = "Category"
hdr[1].text = "Items"
hdr[2].text = "Complexity"

rows = [
    ("No Issues — Just Do It", "6", "Quick field renames, layout changes"),
    ("Watch Out — Needs Discussion", "6", "Design decisions before building"),
    ("SiteTracker Fields", "14+", "Sync script expansion needed"),
    ("Sub-Status Logic", "4 stages", "Dependent picklist design"),
]
for cat, items, notes in rows:
    row = table.add_row().cells
    row[0].text = cat
    row[1].text = items
    row[2].text = notes

doc.add_paragraph("")

# ============================================================
# EMAIL 1: Opportunity Page
# ============================================================
doc.add_heading("Email 1: Opportunity Page Changes", level=2)

# --- Stages ---
doc.add_heading("Stages", level=3)

doc.add_paragraph("Taylor's Questions:", style='List Bullet')
items = [
    "How do sub-statuses work under Prospecting and Under Contract?",
    "How do On Hold / Closed stages capture a reason?",
    'Is "Closed" = signed/handed off to engineering, or closed lost?',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet 2')

doc.add_heading("Assessment:", level=4)
p = doc.add_paragraph()
p.add_run("Recommendation: Dependent picklist.").bold = True
doc.add_paragraph(
    "Salesforce doesn't natively support nested picklists/sub-stages. The standard approach is a "
    "dependent picklist — a Sub_Status__c field whose available values change based on the current "
    "StageName. This is built-in Salesforce functionality and works well."
)
doc.add_paragraph("Options considered:")
doc.add_paragraph("Dependent picklist (Sub_Status__c controlled by StageName) — RECOMMENDED", style='List Bullet')
doc.add_paragraph("Path with guidance — visual only, doesn't enforce sub-status tracking", style='List Bullet')
doc.add_paragraph("")
doc.add_paragraph(
    "Action needed: Define the exact sub-status values for each stage with Taylor before building. "
    "We already discussed 'Design Input Needed' and 'Submit to Engineering' under Under Contract. "
    "Need the full list for Prospecting, On Hold, and Closed."
)

doc.add_paragraph("")

# --- Opportunity Information ---
doc.add_heading("Opportunity Information", level=3)

# Agreement Name -> Site Name
doc.add_heading('1. Rename "Agreement Name" → "Site Name"', level=4)
p = doc.add_paragraph()
r = p.add_run("✓ No issues. ")
r.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
p.add_run("Makes sense — avoids confusion with IronClad/Agreement records. Simple field label change.")

# Lock Opportunity Name
doc.add_heading("2. Lock Opportunity Name after creation", level=4)
p = doc.add_paragraph()
r = p.add_run("⚠ Watch out. ")
r.font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
p.add_run(
    "Taylor's reasoning is valid — name changes cascade to other systems (SiteTracker, Monday.com, IronClad). "
    "But a validation rule that fully blocks edits is too rigid. The first time someone has a typo, they'll be stuck."
)
doc.add_paragraph("Better approach options:", style='List Bullet')
doc.add_paragraph(
    "Validation rule that allows edits only for System Admin profile (or a custom permission)", style='List Bullet 2'
)
doc.add_paragraph(
    "Flow that allows the edit but logs the change and sends a notification to Taylor/admin", style='List Bullet 2'
)
doc.add_paragraph(
    "Approval process — user requests name change, admin approves (heaviest option)", style='List Bullet 2'
)
doc.add_paragraph("")
doc.add_paragraph("Action needed: Decide which approach with Taylor. Recommend option 1 (profile-restricted).")

# Close Date
doc.add_heading("3. Close Date semantics", level=4)
p = doc.add_paragraph()
r = p.add_run("⚠ Watch out. ")
r.font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
p.add_run(
    "Taylor asked if Close Date serves as projected close OR PAL signed date. "
    "Making one field serve two purposes will cause reporting headaches — you can never compare projected vs actual."
)
doc.add_paragraph("Recommendation: Two fields:", style='List Bullet')
doc.add_paragraph("Projected_Close_Date__c (already exists) — BDM sets this during discussions", style='List Bullet 2')
doc.add_paragraph(
    "Standard CloseDate — actual close/PAL signed date (set when deal is won)", style='List Bullet 2'
)
doc.add_paragraph("Action needed: Confirm with Taylor that two fields is acceptable.")

# Amount -> Door Fee
doc.add_heading('4. Repurpose "Amount" as "Door Fee"', level=4)
p = doc.add_paragraph()
r = p.add_run("⚠ Watch out. ")
r.font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
p.add_run(
    "Amount is a standard Salesforce field tied to forecasting, pipeline reports, and out-of-box dashboards. "
    "Repurposing it will break any reporting that references Amount."
)
doc.add_paragraph(
    "Recommendation: Leave Amount hidden and create a new Door_Fee__c currency field. Costs nothing and avoids conflicts.",
    style='List Bullet'
)

doc.add_paragraph("")

# --- Property Details ---
doc.add_heading("Property Details", level=3)

doc.add_heading("5. Fix property type labels (mixed up)", level=4)
p = doc.add_paragraph()
r = p.add_run("✓ No issues. ")
r.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
p.add_run("Taylor's proposed structure is clean and correct:")

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Light Grid Accent 1'
hdr = table2.rows[0].cells
hdr[0].text = "Field"
hdr[1].text = "Type"
hdr[2].text = "Values"

field_rows = [
    ("Build Type", "Picklist", "FTTU, FTTB"),
    ("Brownfield/Greenfield", "Picklist or Checkbox", "Brownfield/Greenfield — or 'New Construction?' checkbox"),
    ("Category", "Picklist", "Cat 1, Cat 2, Cat 3"),
    ("MDU or SFU", "Picklist", "MDU, SFU (remove MHP — redundant with SFU)"),
    ("Property Type", "Picklist", "Apartments, Condos, Townhomes, Private SFU Neighborhood, "
     "Single Family Rental Homes, Mixed Use, Manufactured/Mobile Homes, "
     "Senior/Assisted Living, Commercial/Business (remove 'MDU')"),
]
for field, ftype, values in field_rows:
    row = table2.add_row().cells
    row[0].text = field
    row[1].text = ftype
    row[2].text = values

doc.add_paragraph("")

# HOA
doc.add_heading("6. HOA handling", level=4)
p = doc.add_paragraph()
r = p.add_run("⚠ Watch out. ")
r.font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
p.add_run(
    'Taylor suggested prefixing Property Type values with "HOA –" (e.g., "HOA – Condos"). '
    "This is a bad idea — you lose the ability to filter on HOA independently from property type."
)
doc.add_paragraph(
    "Recommendation: Two separate fields — Property Type picklist + HOA checkbox. Cleaner data, better reporting.",
    style='List Bullet'
)

# Units
doc.add_heading('7. Rename "Units" → "Living Units"', level=4)
p = doc.add_paragraph()
r = p.add_run("✓ No issues. ")
r.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
p.add_run("Simple label change. Also removes need for the info icon.")

# City/State/Zip
doc.add_heading("8. City/State/Zip — keep or remove?", level=4)
p = doc.add_paragraph()
r = p.add_run("⚠ Keep them. ")
r.font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
p.add_run(
    "Taylor asked if separate City/State/Zip fields are necessary. They are — essential for geographic "
    "filtering, market-level reporting, and the SiteTracker sync uses them. Don't remove. "
    "Still worth checking with Jake/Niraj to confirm sync requirements."
)

doc.add_paragraph("")

# --- ISP Information ---
doc.add_heading("ISP Information", level=3)

doc.add_heading("9. Multi-select ISP picklists", level=4)
p = doc.add_paragraph()
r = p.add_run("⚠ Watch out. ")
r.font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
p.add_run(
    "Multi-select fields in Salesforce are notoriously painful for reporting. SOQL can't use '=' on them — "
    "only INCLUDES(). They don't work well in dashboard filters or list view filters. "
    "Taylor specifically said she wants to filter and report by ISP, which is exactly where multi-selects struggle."
)
doc.add_paragraph("Better alternatives:", style='List Bullet')
doc.add_paragraph(
    "Two single-select fields: Primary ISP + Secondary ISP — covers Taylor's 'sometimes two ISPs' scenario",
    style='List Bullet 2'
)
doc.add_paragraph(
    "Junction object (ISP Assignment) — most flexible but heaviest to build, overkill for 5 ISP values",
    style='List Bullet 2'
)
doc.add_paragraph("")
doc.add_paragraph(
    "Recommendation: Primary ISP + Secondary ISP single-select picklists with values: "
    "AT&T, Atlas, FiberFirst, Lumen/Quantum Fiber, Ting. "
    "Discuss with Taylor."
)

# Incumbent
doc.add_heading("10. Move Incumbent fields into ISP section", level=4)
p = doc.add_paragraph()
r = p.add_run("✓ No issues. ")
r.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
p.add_run("Better organization. Three fields:")
doc.add_paragraph("Incumbent Provider (text or picklist)", style='List Bullet')
doc.add_paragraph("Incumbent Agreement Type (picklist: EMA, NEMA, Bulk)", style='List Bullet')
doc.add_paragraph("Incumbent Agreement Expiration Date (date)", style='List Bullet')

doc.add_paragraph("")

# ============================================================
# EMAIL 2: SiteTracker Project Details
# ============================================================
doc.add_heading("Email 2: SiteTracker Project Detail View", level=2)

doc.add_heading("Remove from view", level=3)
remove_items = [
    "Duplicate Project Number (shown twice)",
    "Monday.com Name (obsolete post-migration)",
    "City",
    "State",
    "PAL Signed Date",
    '"In MDU Opportunities" checkbox',
]
for item in remove_items:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
r = p.add_run("✓ No issues. ")
r.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
p.add_run("All straightforward layout changes.")

doc.add_heading("Add to view", level=3)
add_items = [
    "Total Units",
    "Total # of Living Units",
    "Reason for Hold",
    "Confirmed w/ Client Eng Walk Date",
    "Eng Site Walk (A)",
    "Design (1st Draft) Complete based on (A)",
    "Design Phase Complete (A)",
    "Submit the design to the Client (A)",
    "Complete PreCon walk with GC (A)",
    "MDU Construction Start (F)",
    "MDU Construction Start (A)",
    "MDU Construction Complete (F)",
    "MDU Construction Complete (A)",
    '2nd ISP Activation (A) — Taylor unsure of exact SiteTracker field name',
]
for item in add_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph("")
doc.add_heading("Assessment:", level=4)
p = doc.add_paragraph()
r = p.add_run("⚠ Biggest work item. ")
r.font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
p.add_run("Two things need to happen before these can show up in Salesforce:")

doc.add_paragraph(
    "1. Check which of these 14 fields exist in the SiteTracker org. The daily sync script "
    "(sync_sitetracker.py) only pulls a subset of SiteTracker fields. Any field Taylor wants to see "
    "in SF needs to be added to the sync query and mapped to a new field on SiteTracker_Project__c.",
    style='List Bullet'
)
doc.add_paragraph(
    '2. "2nd ISP Activation (A)" — need to look up the exact API name in SiteTracker before we can sync it.',
    style='List Bullet'
)
doc.add_paragraph("")
doc.add_paragraph(
    "Action needed: Query SiteTracker org to audit which of these fields exist and get their API names. "
    "Then expand the sync script + create new fields on SiteTracker_Project__c."
)

# Pending discussion
doc.add_heading("Pending: Build Status Labels & Icons", level=3)
doc.add_paragraph(
    "Taylor wants to align Build Status labels and icons with Niraj/Jake so it's clear what each stage means. "
    "This is a meeting, not a build task — schedule it."
)

doc.add_paragraph("")

# ============================================================
# Implementation Order
# ============================================================
doc.add_heading("Suggested Implementation Order", level=2)

doc.add_paragraph(
    "Phase 1 — Quick wins (no discussion needed):", style='List Bullet'
)
phase1 = [
    'Rename "Agreement Name" → "Site Name"',
    'Rename "Units" → "Living Units"',
    "Remove duplicate Project Number from SiteTracker view",
    "Remove Monday.com Name from SiteTracker view",
    "Move Incumbent fields into ISP section",
    "Remove City/State/PAL Signed/In MDU Opps from SiteTracker view",
]
for item in phase1:
    doc.add_paragraph(item, style='List Bullet 2')

doc.add_paragraph(
    "Phase 2 — After Taylor confirms decisions:", style='List Bullet'
)
phase2 = [
    "Create Door_Fee__c field (hide standard Amount)",
    "Fix Property Type picklist structure (5 separate fields)",
    "Add HOA checkbox",
    "Primary ISP + Secondary ISP picklists (or multi-select if Taylor insists)",
    "Incumbent Provider fields (3 new fields)",
    "Close Date: confirm two-field approach",
    "Opportunity Name lock: choose restriction method",
]
for item in phase2:
    doc.add_paragraph(item, style='List Bullet 2')

doc.add_paragraph(
    "Phase 3 — Requires investigation:", style='List Bullet'
)
phase3 = [
    "Sub-status dependent picklist (need full value list from Taylor)",
    "SiteTracker field expansion (audit ST org, expand sync, add 14 fields)",
    "Build Status labels meeting with Niraj/Jake",
]
for item in phase3:
    doc.add_paragraph(item, style='List Bullet 2')

# Save
doc.save("C:/Users/cass/Work_Projects/SalesForce/taylor-feedback-review.docx")
print("Done")
