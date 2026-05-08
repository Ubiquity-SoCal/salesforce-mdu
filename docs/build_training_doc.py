"""
Generates the MDU Sales Training Guide as a .docx.
Run: python build_training_doc.py
"""

from pathlib import Path
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).parent / "MDU Sales Training Guide.docx"
DATA_MODEL_PNG = Path(__file__).parent / "_diagram_data_model.png"
LIFECYCLE_PNG = Path(__file__).parent / "_diagram_lifecycle.png"

NAVY_HEX = "#1F3A68"
ACCENT_HEX = "#2E7D32"
GRAY_HEX = "#555555"
SOFT_HEX = "#E8EEF7"
WARN_HEX = "#B71C1C"
WARN_SOFT_HEX = "#FCE4EC"

NAVY = RGBColor(0x1F, 0x3A, 0x68)
GRAY = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0x2E, 0x7D, 0x32)


def set_run(run, *, bold=False, size=None, color=None, italic=False):
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def add_para(doc, text, *, style=None, bold=False, size=None, color=None, italic=False, after=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, bold=bold, size=size, color=color, italic=italic)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    return p


def add_h1(doc, text):
    p = doc.add_heading(level=1)
    run = p.add_run(text)
    run.font.size = Pt(20)
    run.font.color.rgb = NAVY
    run.bold = True


def add_h2(doc, text):
    p = doc.add_heading(level=2)
    run = p.add_run(text)
    run.font.size = Pt(15)
    run.font.color.rgb = NAVY
    run.bold = True


def add_h3(doc, text):
    p = doc.add_heading(level=3)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = ACCENT
    run.bold = True


def add_bullets(doc, items):
    for item in items:
        if isinstance(item, tuple):
            label, body = item
            p = doc.add_paragraph(style="List Bullet")
            r1 = p.add_run(f"{label}: ")
            r1.bold = True
            p.add_run(body)
        else:
            doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_callout(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r1 = p.add_run(f"{label}  ")
    r1.bold = True
    r1.font.color.rgb = ACCENT
    r2 = p.add_run(text)
    r2.italic = True
    r2.font.color.rgb = GRAY


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = NAVY
    for r_i, row in enumerate(rows, start=1):
        for c_i, val in enumerate(row):
            table.rows[r_i].cells[c_i].text = str(val)
    doc.add_paragraph()


def _box(ax, x, y, w, h, label, *, fill=SOFT_HEX, edge=NAVY_HEX, text=NAVY_HEX, fontsize=11, bold=True):
    box = FancyBboxPatch(
        (x, y), w, h,
        boxstyle="round,pad=0.02,rounding_size=0.08",
        linewidth=1.6, edgecolor=edge, facecolor=fill,
    )
    ax.add_patch(box)
    ax.text(x + w / 2, y + h / 2, label,
            ha="center", va="center",
            fontsize=fontsize, color=text,
            fontweight="bold" if bold else "normal",
            wrap=True)


def _arrow(ax, x1, y1, x2, y2, *, color=GRAY_HEX, label=None, label_offset=(0, 0.12), style="-|>", lw=1.4):
    arrow = FancyArrowPatch(
        (x1, y1), (x2, y2),
        arrowstyle=style, mutation_scale=14,
        color=color, linewidth=lw,
    )
    ax.add_patch(arrow)
    if label:
        mx, my = (x1 + x2) / 2 + label_offset[0], (y1 + y2) / 2 + label_offset[1]
        ax.text(mx, my, label, ha="center", va="center",
                fontsize=8.5, color=GRAY_HEX, style="italic")


def build_data_model_diagram():
    fig, ax = plt.subplots(figsize=(11, 7.5))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 10)
    ax.axis("off")

    ax.text(7, 9.55, "MDU Sales Data Model",
            ha="center", va="center", fontsize=16, fontweight="bold", color=NAVY_HEX)
    ax.text(7, 9.15, "How the records relate. Opportunity is the center.",
            ha="center", va="center", fontsize=10, color=GRAY_HEX, style="italic")

    # Account (top center-left)
    _box(ax, 1.0, 7.3, 2.6, 1.1, "ACCOUNT\nProperty mgmt co.\nor owner entity")

    # Property Location (top right)
    _box(ax, 10.4, 7.3, 2.6, 1.1, "PROPERTY LOCATION\nThe serviceable\nproperty / address")

    # Property Unit (far right, lower)
    _box(ax, 10.4, 5.4, 2.6, 1.0, "PROPERTY UNIT\nIndividual units\n(MDU: passive)", fontsize=10)

    # Opportunity (center, bigger and accented)
    _box(ax, 5.5, 5.5, 3.0, 1.5, "OPPORTUNITY\nThe property pursuit\n(MDU record type)",
         fill="#FFF8E1", edge=ACCENT_HEX, text=ACCENT_HEX, fontsize=12)

    # Contact (left middle)
    _box(ax, 0.4, 4.3, 2.4, 1.1, "CONTACT\nA person\n(PM, owner, HOA)")

    # Opportunity Contact junction (between Contact and Opp)
    _box(ax, 3.2, 4.5, 2.0, 0.9,
         "OPPORTUNITY\nCONTACT (junction)\nrole on this Opp",
         fill="#F1F8E9", edge=ACCENT_HEX, text=ACCENT_HEX, fontsize=8.5)

    # Agreement (below Opp)
    _box(ax, 5.3, 2.6, 1.7, 1.0, "AGREEMENT\nPAL / ROE / EMA\n/ Bulk / NEMA", fontsize=9)

    # Note (below Opp center)
    _box(ax, 7.2, 2.6, 1.4, 1.0, "NOTE\nEvery convo,\ndecision, update", fontsize=9)

    # SiteTracker Project (bottom right)
    _box(ax, 10.4, 3.3, 2.6, 1.1,
         "SITETRACKER PROJECT\nLives in separate\nSF org. Build only.",
         fill=WARN_SOFT_HEX, edge=WARN_HEX, text=WARN_HEX, fontsize=9)

    # ----- Arrows -----
    # Account -> Opportunity
    _arrow(ax, 2.3, 7.3, 6.6, 7.0, label="owns many")
    # Account -> Property Location
    _arrow(ax, 3.6, 7.85, 10.4, 7.85, label="has many")
    # Property Location -> Opportunity
    _arrow(ax, 10.4, 7.5, 8.5, 6.6, label="linked to")
    # Property Location -> Property Unit
    _arrow(ax, 11.7, 7.3, 11.7, 6.4, label="has many")
    # Contact -> Opportunity Contact
    _arrow(ax, 2.8, 4.95, 3.2, 4.95, label=None)
    # Opportunity Contact -> Opportunity
    _arrow(ax, 5.2, 5.0, 5.5, 5.7, label=None)
    # Opportunity -> Agreement
    _arrow(ax, 6.5, 5.5, 6.1, 3.6, label="has many", label_offset=(-0.3, 0))
    # Opportunity -> Note
    _arrow(ax, 7.4, 5.5, 7.8, 3.6, label="has many", label_offset=(0.35, 0))
    # Opportunity -> SiteTracker (dashed-ish via lighter color)
    _arrow(ax, 8.5, 6.0, 10.4, 4.1, color=WARN_HEX, label="when build starts", label_offset=(0, 0.2))

    # Legend
    ax.text(0.4, 1.5, "Legend:", fontsize=10, fontweight="bold", color=NAVY_HEX)
    legend_items = [
        ("Standard Salesforce object", SOFT_HEX, NAVY_HEX),
        ("Opportunity (your home base)", "#FFF8E1", ACCENT_HEX),
        ("Junction object (linking record)", "#F1F8E9", ACCENT_HEX),
        ("External system handoff", WARN_SOFT_HEX, WARN_HEX),
    ]
    for i, (text, fill, edge) in enumerate(legend_items):
        x_pos = 0.4 + (i * 3.4)
        _box(ax, x_pos, 0.65, 0.45, 0.45, "", fill=fill, edge=edge)
        ax.text(x_pos + 0.6, 0.88, text, fontsize=8.5, va="center", color=GRAY_HEX)

    plt.tight_layout()
    plt.savefig(DATA_MODEL_PNG, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def build_lifecycle_diagram():
    fig, ax = plt.subplots(figsize=(11, 7.0))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.axis("off")

    ax.text(7, 8.55, "Opportunity Lifecycle",
            ha="center", va="center", fontsize=16, fontweight="bold", color=NAVY_HEX)
    ax.text(7, 8.15, "Stages an MDU pursuit moves through, plus the off-ramps.",
            ha="center", va="center", fontsize=10, color=GRAY_HEX, style="italic")

    # Main path stages
    stages = [
        ("PROSPECTING\n10%", 0.5),
        ("ENGAGED\n30%", 3.1),
        ("CONTRACT\nNEGOTIATIONS\n50%", 5.7),
        ("ROE SECURED\n50%", 8.3),
        ("UNDER CONTRACT\n100%", 10.9),
    ]
    box_w, box_h = 2.3, 1.1
    y_main = 5.5

    for i, (label, x) in enumerate(stages):
        is_last = (i == len(stages) - 1)
        fill = "#E8F5E9" if is_last else SOFT_HEX
        edge = ACCENT_HEX if is_last else NAVY_HEX
        text = ACCENT_HEX if is_last else NAVY_HEX
        _box(ax, x, y_main, box_w, box_h, label, fill=fill, edge=edge, text=text, fontsize=10)

    # Arrows between stages
    for i in range(len(stages) - 1):
        x_from = stages[i][1] + box_w
        x_to = stages[i + 1][1]
        _arrow(ax, x_from, y_main + box_h / 2, x_to, y_main + box_h / 2)

    # Required-field annotations under each stage
    annotations = [
        (0.5, "Sales Status\nrequired"),
        (3.1, "Property contact\ncaptured"),
        (5.7, "Agreement\nin Create or Review"),
        (8.3, "ROE Agreement\nmarked Signed"),
        (10.9, "PAL Signed +\nSiteTracker linked"),
    ]
    for x, txt in annotations:
        ax.text(x + box_w / 2, y_main - 0.3, txt,
                ha="center", va="top", fontsize=8.5, color=GRAY_HEX, style="italic")

    # SiteTracker handoff arrow
    _arrow(ax, 11.9, y_main, 12.8, 3.2, color=WARN_HEX,
           label="Build owned\nby SiteTracker", label_offset=(0.6, 0))
    _box(ax, 11.4, 2.0, 2.2, 1.2,
         "SITETRACKER\nORG\n(construction)",
         fill=WARN_SOFT_HEX, edge=WARN_HEX, text=WARN_HEX, fontsize=9)

    # Off-ramps: On Hold and Closed Lost
    _box(ax, 3.5, 1.8, 2.3, 1.0, "ON HOLD\nHold Reason\nrequired",
         fill="#FFF3E0", edge="#E65100", text="#E65100", fontsize=9.5)
    _box(ax, 7.0, 1.8, 2.3, 1.0, "CLOSED LOST\nLoss Reason\nrequired",
         fill=WARN_SOFT_HEX, edge=WARN_HEX, text=WARN_HEX, fontsize=9.5)

    # Off-ramp arrows: from anywhere on the main path
    ax.text(4.6, 3.3, "from any\nactive stage", ha="center", fontsize=8, color=GRAY_HEX, style="italic")
    _arrow(ax, 4.6, 5.4, 4.6, 2.85, color="#E65100")
    _arrow(ax, 8.2, 5.4, 8.2, 2.85, color=WARN_HEX)

    # Legend / notes
    ax.text(0.4, 0.95, "Notes:", fontsize=10, fontweight="bold", color=NAVY_HEX)
    notes = [
        "Stages are how leadership reads pipeline. Move them as reality changes.",
        "On Hold is reversible. Closed Lost ends the pursuit (99% of losses are existing fiber).",
        "Once Under Contract, build progress lives in SiteTracker, not on the Opportunity.",
    ]
    for i, n in enumerate(notes):
        ax.text(0.4, 0.6 - i * 0.22, "  " + n, fontsize=8.8, color=GRAY_HEX, va="top")

    plt.tight_layout()
    plt.savefig(LIFECYCLE_PNG, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def build():
    build_data_model_diagram()
    build_lifecycle_diagram()
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---------------- Title page ----------------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("MDU Sales Training Guide")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Salesforce: Day-to-Day for the MDU Team")
    r.font.size = Pt(14)
    r.font.color.rgb = GRAY

    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = intro.add_run(
        "This guide walks the MDU sales team through how we work in Salesforce now that we have moved off Monday.com. "
        "Each section maps to something you will do during the week. Use it as a follow-along during training and as a desk reference after."
    )
    r.font.size = Pt(11)
    r.italic = True
    r.font.color.rgb = GRAY

    doc.add_page_break()

    # ---------------- Table of contents ----------------
    add_h1(doc, "What this guide covers")
    add_numbered(doc, [
        "Why we moved to Salesforce",
        "Logging in and getting around the MDU Sales app",
        "How the data fits together (the model in plain English)",
        "Opportunity stages and what each one means",
        "Daily workflows (step by step)",
        "Agreements: PAL, ROE, EMA, Bulk and beyond",
        "The MDU Tracker (your main view)",
        "Reports and dashboards",
        "How SiteTracker fits in",
        "Common gotchas and validation rules",
        "What is coming next",
    ])

    doc.add_page_break()

    # ============== 1. WHY ==============
    add_h1(doc, "1. Why we moved to Salesforce")
    add_para(doc,
        "Monday.com served us well as a list of properties, but it was a list, not a system. "
        "Salesforce gives us one place where the property pursuit, the people involved, the agreements, the build status, "
        "and eventually the billing all live together and stay in sync."
    )
    add_h3(doc, "What changes for you")
    add_bullets(doc, [
        ("One record per property", "the Opportunity is the property pursuit. Everything else hangs off it."),
        ("Notes are first class", "every conversation, decision, and update goes on the Opportunity as a Note. No more hunting through Monday updates or email."),
        ("Agreements are tracked separately", "PAL, ROE, EMA, Bulk and so on each get their own record under the Opportunity, with their own status."),
        ("Stages are the source of truth", "moving an Opportunity through the stages is how leadership sees pipeline. Keep them current."),
        ("Build tracking still lives in SiteTracker", "the Opportunity does not track construction. Once an agreement is signed, SiteTracker takes over."),
    ])
    add_h3(doc, "What was migrated")
    add_bullets(doc, [
        "3,229 MDU Opportunities from Monday.com",
        "26 Accounts and 128 Contacts",
        "1,291 Agreements (status carried over from Monday columns)",
        "13,269 historical Notes (every Monday update preserved by author and date)",
    ])

    doc.add_page_break()

    # ============== 2. LOGIN / NAVIGATION ==============
    add_h1(doc, "2. Logging in and getting around")
    add_h3(doc, "Where to log in")
    add_bullets(doc, [
        ("Lightning URL", "https://fun-power-747.lightning.force.com"),
        ("Username", "your @ubiquitygp.com Salesforce username"),
        ("If you cannot log in", "use Forgot Password on the login page, or ping Cass."),
    ])

    add_h3(doc, "Switching to the MDU Sales app")
    add_bullets(doc, [
        "Top left, click the App Launcher (the nine-dot grid).",
        "Type \"MDU Sales\" and pin it. This is the app you live in day to day.",
        "If you also work Business pursuits, switch to the Business Sales app from the same launcher.",
    ])

    add_h3(doc, "Tabs you will use")
    add_table(doc,
        headers=["Tab", "What it is", "When to use it"],
        rows=[
            ["Opportunities", "Every property pursuit", "Most of your day. Open one to update stage, log notes, manage agreements."],
            ["Accounts", "Property management companies and ownership entities", "When you need to see all properties tied to one company."],
            ["Contacts", "People (PMs, owners, leasing, HOA, GCs)", "Search for a person, see what properties they touch."],
            ["Agreements", "PAL, ROE, EMA, Bulk records", "Pipeline view of contracts in flight."],
            ["Property Locations", "Serviceable properties", "Address-level record. One Account can have many Locations."],
            ["Reports", "Saved reports", "Pull lists for meetings or standups."],
            ["Dashboards", "Visual rollups", "Open the MDU Sales Dashboard for a snapshot of the pipeline."],
        ],
        col_widths=[1.4, 2.0, 2.8],
    )

    doc.add_page_break()

    # ============== 3. DATA MODEL ==============
    add_h1(doc, "3. How the data fits together")
    add_para(doc,
        "If you remember nothing else, remember this: the Opportunity is the center of the universe. "
        "Everything else hangs off it."
    )

    add_h3(doc, "The pieces")
    add_bullets(doc, [
        ("Account", "the property management company or ownership entity (e.g., Greystar, a single-property LLC)."),
        ("Contact", "an actual person (Property Manager, Owner, HOA, Leasing Contact, GC, Legal, Broker)."),
        ("Opportunity", "the property pursuit. One per property. This is where your work lives."),
        ("Opportunity Contact", "links a Contact to an Opportunity with a Role. A Contact can be on many Opps; an Opp can have many Contacts."),
        ("Agreement", "a child of the Opportunity. PAL, ROE, EMA, Bulk, etc. One Opp can have many Agreements over its life."),
        ("Property Location", "the serviceable property record. Linked to the Opportunity. Address, unit count, parcel data live here."),
        ("Property Unit", "an individual unit within a Location. Mostly used by Business; passive for MDU."),
        ("SiteTracker Project", "the build record in our separate construction Salesforce org. Linked to the Opportunity once construction starts."),
    ])

    add_h3(doc, "How they connect")
    add_para(doc, "Read this top to bottom:")
    add_bullets(doc, [
        "Account (Greystar) has many Property Locations and many Opportunities.",
        "Each Opportunity points to one Property Location and is tied to one Account.",
        "Each Opportunity has many Contacts via Opportunity Contact (each with a Role).",
        "Each Opportunity has many Agreements (PAL, ROE, etc.).",
        "Each Opportunity has many Notes (chronological, every conversation).",
        "Once construction begins, the Opportunity links out to its SiteTracker Project.",
    ])

    add_callout(doc, "Why a junction object for Contacts?",
        "A Property Manager often manages dozens of properties. We need that one Contact record connected to every "
        "Opportunity she touches, with the role on each. Opportunity Contact is what makes that work."
    )

    doc.add_page_break()

    # ============== 4. STAGES ==============
    add_h1(doc, "4. Opportunity stages and what each means")
    add_para(doc,
        "Stages tell the company where a property is in the pursuit. They are how leadership reads the pipeline, "
        "so keeping them honest is the single most important habit."
    )

    add_table(doc,
        headers=["Stage", "Probability", "What it means", "What is required"],
        rows=[
            ["Prospecting", "10%", "We have the property on our radar. Outreach pending or in progress.", "Sales Status: Contact Pending or Reached Out, Pending Response."],
            ["Engaged", "30%", "We are in conversation with the property. Real dialogue happening.", "Property contact info captured."],
            ["Contract Negotiations", "50%", "Terms being discussed. Agreement language going back and forth.", "At least one Agreement record (Status: Create or Review)."],
            ["ROE Secured", "50%", "Right of Entry locked, paving the way for build.", "ROE Agreement marked Signed."],
            ["Under Contract", "100%", "PAL signed. SiteTracker takes over for build.", "PAL Agreement marked Signed. SiteTracker Project linked."],
            ["On Hold", "0%", "Pursuit paused for a defined reason.", "Hold Reason required."],
            ["Closed Lost", "0%", "Deal will not happen. Vast majority are existing fiber.", "Loss Reason required."],
        ],
        col_widths=[1.4, 0.8, 2.4, 1.9],
    )

    add_h3(doc, "Sales Status (used during Prospecting)")
    add_bullets(doc, [
        ("Contact Pending", "we do not yet have property contact info."),
        ("Reached Out, Pending Response", "outreach made, waiting on a reply."),
        ("Research Completed", "background work done, ready to engage."),
    ])

    add_h3(doc, "Hold Reasons (required when Stage = On Hold)")
    add_bullets(doc, [
        "Ownership Change",
        "Budget or Timing",
        "Pending Legal Review",
        "Market Conditions",
        "Other",
    ])

    add_h3(doc, "Loss Reasons (required when Stage = Closed Lost)")
    add_bullets(doc, [
        "Existing Fiber (most common)",
        "Existing Contract",
        "Not Interested",
        "No Contact Info",
        "Chose Another Provider",
        "Unserviceable",
        "Other",
    ])

    doc.add_page_break()

    # ============== 5. WORKFLOWS ==============
    add_h1(doc, "5. Daily workflows, step by step")

    add_h2(doc, "5.1 Create a new Opportunity")
    add_numbered(doc, [
        "Opportunities tab, then click New.",
        "Pick the MDU record type when prompted.",
        "Fill the basics: Opportunity Name (use the property name), Account, Stage, Projected Close Date.",
        "Property Details section: Units, Property Type, City, State, Zip (these are required, you cannot save without them).",
        "If Stage is Prospecting, set a Sales Status.",
        "Save. The record opens, and you can start logging notes and adding contacts.",
    ])
    add_callout(doc, "Naming convention",
        "Property name first, units second when ambiguity exists, e.g. \"Olympus Waterford 25\". Avoid leading prefixes like \"MDU - \"."
    )

    add_h2(doc, "5.2 Log a note on an Opportunity")
    add_numbered(doc, [
        "Open the Opportunity.",
        "Find the Notes related list (left sidebar on the record page).",
        "Click New, give it a short title, write the note. Use this for every meaningful conversation, decision, or status update.",
        "Save. The newest note shows on top.",
    ])
    add_callout(doc, "Note discipline",
        "If a future you (or a teammate) opens this Opportunity in six months, the Notes should tell the story. "
        "Date and topic in the title beats vague titles like \"update\"."
    )

    add_h2(doc, "5.3 Add a Contact to an Opportunity")
    add_numbered(doc, [
        "Open the Opportunity.",
        "In the Contacts related list (left sidebar), click New.",
        "Search for the Contact. If they exist (Property Manager that already touches other Opps), pick the existing record. Do not duplicate people.",
        "Set their Role: Property Manager, Owner, Leasing Contact, HOA Contact, GC, Developer, Legal, Broker, or Other.",
        "Save.",
    ])
    add_callout(doc, "If the Contact does not exist yet",
        "Create it from the Contacts tab first (with Account if known), then come back to the Opportunity and link them. "
        "This keeps the Contact reusable across every property they touch."
    )

    add_h2(doc, "5.4 Create an Agreement")
    add_numbered(doc, [
        "Open the Opportunity.",
        "Agreements related list (left sidebar), click New.",
        "Pick the Agreement Type: PAL, ROE, EMA, Bulk, NEMA, PAL Addendum, MSA Addendum, 2nd ISP NEMA, 2nd ISP MSA Addendum.",
        "Set Status (start at Create, then Review, then Sign as it progresses).",
        "Set Requested Date when it goes out.",
        "Set Signer (Contact lookup) and Signed Date when it lands.",
        "If you already have an IronClad ID, paste it into IronClad ID. Otherwise leave blank for now.",
        "Save.",
    ])
    add_callout(doc, "PAL is the gateway",
        "Nothing real happens on a property until the PAL is signed. Other agreements (Bulk, EMA) often run in parallel but PAL is what unlocks engineering."
    )

    add_h2(doc, "5.5 Move an Opportunity to Under Contract")
    add_numbered(doc, [
        "PAL Agreement Status set to Sign or Completed with Signed Date filled.",
        "On the Opportunity, change Stage to Under Contract.",
        "Salesforce will auto-set probability to 100%.",
        "From here, build progress is tracked in SiteTracker. The Opp does not change stage again unless something goes sideways.",
    ])

    add_h2(doc, "5.6 Put an Opportunity On Hold")
    add_numbered(doc, [
        "Open the Opportunity.",
        "Change Stage to On Hold.",
        "Set Hold Reason (required, Salesforce will not let you save without it).",
        "Add a Note explaining context: who paused it, when to revisit.",
        "Save.",
    ])

    add_h2(doc, "5.7 Close an Opportunity as Lost")
    add_numbered(doc, [
        "Open the Opportunity.",
        "Change Stage to Closed Lost.",
        "Set Loss Reason (required).",
        "Add a Note with detail: why, who told us, any follow-up worth tracking.",
        "Save.",
    ])

    doc.add_page_break()

    # ============== 6. AGREEMENTS ==============
    add_h1(doc, "6. Agreements")
    add_para(doc,
        "Each Opportunity has its own track of Agreements. Most properties end up with several by the time activation happens. "
        "Knowing which type does what is half the job."
    )

    add_h3(doc, "Agreement types")
    add_table(doc,
        headers=["Type", "What it does"],
        rows=[
            ["PAL", "Property Access License. The gateway agreement. Required before engineering."],
            ["ROE", "Right of Entry. Covers physical access to the property for survey and build."],
            ["EMA", "Easement / marketing access. Often parallel to PAL."],
            ["Bulk", "Bulk service agreement when the property is paying for all units."],
            ["NEMA", "Non-exclusive marketing agreement."],
            ["PAL Addendum", "Modification to an existing PAL."],
            ["MSA Addendum", "Modification to a master services agreement."],
            ["2nd ISP NEMA", "When a second ISP is also marketing on the property."],
            ["2nd ISP MSA Addendum", "Companion to the above."],
        ],
        col_widths=[1.8, 4.7],
    )

    add_h3(doc, "Status flow (aligned to IronClad)")
    add_bullets(doc, [
        ("Create", "drafting in progress on our side."),
        ("Review", "out for internal or counterparty review."),
        ("Sign", "signature collection underway."),
        ("Completed", "fully executed."),
        ("Archive", "completed and filed for the record."),
        ("Paused", "stalled, waiting on something external."),
        ("Cancelled", "will not happen."),
    ])

    add_h3(doc, "IronClad ID")
    add_para(doc,
        "Once IronClad integration lands, every signed agreement should have an IronClad ID. Today, populate it when you have it. "
        "Going forward, Agreements with no IronClad ID will show on a cleanup report so we can backfill. "
        "Treat the IronClad ID as the audit trail."
    )

    doc.add_page_break()

    # ============== 7. TRACKER ==============
    add_h1(doc, "7. The MDU Tracker")
    add_para(doc,
        "The Tracker is the single most useful view in the app. It is the home page experience for the MDU Sales app: "
        "filtered, paginated, summarized, and built around how we actually work."
    )

    add_h3(doc, "What it shows")
    add_bullets(doc, [
        "Property name, Stage, SiteTracker Build Status, Owner, RE Assigned, Projected Close Date.",
        "Summary bar at the top: Opps in view, total Units, completed Agreements.",
        "50 Opportunities per page with Load More at the bottom.",
    ])

    add_h3(doc, "Filters")
    add_bullets(doc, [
        ("Owner", "filter to your own pipeline or someone else's."),
        ("RE Assigned", "filter by the RE owner (TF, JB, RS, etc.)."),
        ("Stage", "MDU stages only, scoped to this app."),
        ("Date range", "by Projected Close Date."),
    ])

    add_h3(doc, "How to use it")
    add_bullets(doc, [
        "Mondays: filter to your Owner, sort by Stage, walk every Opp in Engaged or Contract Negotiations.",
        "Mid-week: scan On Hold, look for any that should re-engage.",
        "Fridays: anything Closing this month with a stale Stage gets a Note and a status check.",
    ])

    doc.add_page_break()

    # ============== 8. REPORTS ==============
    add_h1(doc, "8. Reports and dashboards")

    add_h3(doc, "MDU Sales Dashboard")
    add_para(doc,
        "Six charts, filterable by Opportunity Owner. Open it from the Dashboards tab. Use this for your own pipeline reviews "
        "and as the snapshot you bring to standups."
    )
    add_bullets(doc, [
        "Opportunities by Stage",
        "Opportunities by Owner",
        "Opportunities by State",
        "Units by Stage",
        "Stage by Owner Matrix",
        "SiteTracker Coverage",
    ])

    add_h3(doc, "Saved reports (in MDU Sales Reports folder)")
    add_bullets(doc, [
        "MDU Opportunities by Stage",
        "MDU Opportunities by Owner",
        "MDU Opportunities by State",
        "MDU Units by Stage",
        "Stage by Owner Matrix",
        "SiteTracker Coverage",
    ])

    add_h3(doc, "Campaign Dashboard (for tagged initiatives)")
    add_para(doc,
        "When a set of properties is part of a named initiative (for example the 9-25 MDU ROE Project), they are tagged via Campaign. "
        "The Campaign Dashboard shows pipeline progress for just that group of Opportunities."
    )

    doc.add_page_break()

    # ============== 9. SITETRACKER ==============
    add_h1(doc, "9. How SiteTracker fits in")
    add_para(doc,
        "SiteTracker is a separate Salesforce org used by Engineering and Construction. We do not track build milestones "
        "on the Opportunity in MDU Sales. Once a PAL is signed, the property crosses over to SiteTracker for the build."
    )

    add_h3(doc, "What you see in MDU Sales")
    add_bullets(doc, [
        ("SiteTracker Project ID", "the linking key. Auto-populated when the Project syncs over."),
        ("SiteTracker URL", "click through to the build record in the SiteTracker org."),
        ("ST Build Status", "shown on the Tracker as a virtual column. Reflects the build phase."),
        ("In SiteTracker checkbox", "yes/no flag whether the Opportunity is linked to a SiteTracker Project."),
    ])

    add_h3(doc, "How the sync works")
    add_bullets(doc, [
        "GitHub Action runs every night at 2am Central, syncing SiteTracker Projects into MDU Sales.",
        "Match key is property name and Agreement Name. New projects get linked automatically.",
        "If a property has no link after construction starts, flag it to Cass.",
    ])

    add_callout(doc, "Rule of thumb",
        "If it is about the sale (status, agreements, contacts, conversations) it goes in MDU Sales. "
        "If it is about the build (FDH activation, lit fiber, install schedule) it lives in SiteTracker."
    )

    doc.add_page_break()

    # ============== 10. GOTCHAS ==============
    add_h1(doc, "10. Common gotchas and rules")
    add_para(doc, "These are the validation rules and quirks that catch people their first week.")

    add_h3(doc, "Validation rules you will hit")
    add_bullets(doc, [
        ("City, State, Zip required", "you cannot save an MDU Opportunity without all three. Set them at creation."),
        ("Sales Status required when Stage is Prospecting", "pick Contact Pending, Reached Out Pending Response, or Research Completed."),
        ("Hold Reason required when Stage is On Hold", "pick a reason."),
        ("Loss Reason required when Stage is Closed Lost", "pick a reason."),
        ("Agreement Name is unique across the org", "do not duplicate. If you see a duplicate error, the Opportunity already exists, search first."),
    ])

    add_h3(doc, "Habits to keep")
    add_bullets(doc, [
        "Do not edit the standard Close Date field. Use Projected Close Date for forecasting.",
        "Always reuse an existing Contact before creating a new one. Property Managers wear many hats.",
        "Use the Notes related list for every conversation. The pipeline review meetings depend on it.",
        "Move the Stage when reality changes. Stale stages are worse than empty ones.",
    ])

    add_h3(doc, "If something looks wrong")
    add_bullets(doc, [
        "If you cannot find an Opportunity that should exist, check Closed Lost and On Hold filters first.",
        "If a Contact is on the wrong Account, fix the Contact, not the Opportunity link.",
        "If the SiteTracker link is missing on a built property, send the property name to Cass for the linker script.",
        "If validation rules block a save you think should work, screenshot the error and ping Cass before working around it.",
    ])

    doc.add_page_break()

    # ============== 11. WHAT'S NEXT ==============
    add_h1(doc, "11. What is coming next")
    add_bullets(doc, [
        ("IronClad sync", "agreement status will flow automatically from IronClad once API access lands. Until then, update Status manually as it changes."),
        ("Monday.com retirement", "once the team is comfortable in Salesforce and we are confident the data is clean, Monday.com goes away for MDU."),
        ("ROE and SAQ Excel trackers fold in", "the standalone Excel dashboards become Salesforce reports backed by Agreement records."),
        ("Investor PAL report", "a PAL pipeline report as a tile on the MDU home page, for fundraising reporting."),
        ("Tracker enhancements", "more filters, saved views per user, deeper SiteTracker integration as we go."),
    ])

    # ---------------- Footer / closing ----------------
    doc.add_page_break()
    add_h1(doc, "Quick reference card")
    add_h3(doc, "Required when you change Stage")
    add_table(doc,
        headers=["Stage", "Required field"],
        rows=[
            ["Prospecting", "Sales Status"],
            ["On Hold", "Hold Reason"],
            ["Closed Lost", "Loss Reason"],
            ["Under Contract", "PAL Agreement marked Signed"],
        ],
        col_widths=[2.5, 3.5],
    )

    add_h3(doc, "Where things live")
    add_table(doc,
        headers=["If you want to...", "Go to..."],
        rows=[
            ["See your pipeline at a glance", "MDU Sales Dashboard"],
            ["Walk every active Opp in your name", "MDU Tracker, filter by Owner = you"],
            ["See agreements in flight", "Agreements tab, MDU Agreements list view"],
            ["Find a Property Manager", "Contacts tab, search by name"],
            ["See all properties for a company", "Account record, Opportunities related list"],
            ["Check construction status", "SiteTracker URL on the Opportunity"],
            ["Run a list for a meeting", "Reports tab, MDU Sales Reports folder"],
        ],
        col_widths=[3.0, 3.0],
    )

    add_h3(doc, "Who to ask")
    add_bullets(doc, [
        ("Salesforce questions, validation errors, missing data", "Cass"),
        ("Agreement language and IronClad", "Taylor"),
        ("SiteTracker / build status", "Engineering team"),
        ("Pipeline strategy and stage definitions", "your sales lead"),
    ])

    # ============== Visual Reference ==============
    doc.add_page_break()
    add_h1(doc, "Visual reference")
    add_para(doc,
        "Two diagrams that summarize the whole guide. The first shows how the records relate. "
        "The second shows the stages an Opportunity moves through and the off-ramps."
    )

    add_h2(doc, "Data model: how the records connect")
    add_para(doc,
        "Account is the company. Each Account owns Property Locations (the serviceable property) and Opportunities (the pursuit). "
        "Each Opportunity is tied to one Property Location, has many Contacts via the Opportunity Contact junction, "
        "and accumulates Agreements and Notes. When the build kicks off, the Opportunity hands off to a SiteTracker Project in the separate construction org."
    )
    doc.add_picture(str(DATA_MODEL_PNG), width=Inches(6.8))

    doc.add_paragraph()

    add_h2(doc, "Opportunity lifecycle")
    add_para(doc,
        "Most pursuits follow the path along the top: Prospecting through Under Contract. "
        "Two off-ramps (On Hold and Closed Lost) can happen from any active stage. "
        "Each stage has a small set of required fields, listed under the box."
    )
    doc.add_picture(str(LIFECYCLE_PNG), width=Inches(6.8))

    doc.save(OUT)
    print(f"Wrote: {OUT}")


if __name__ == "__main__":
    build()
